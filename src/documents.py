"""
Загрузка документов и разбиение на чанки для RAG.
Используются только pypdf и python-docx — без лоадеров LangChain, чтобы не тянуть PyTorch.
"""
from pathlib import Path
from typing import List

from langchain_core.documents import Document


def _load_txt(path: Path) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _load_docx(path: Path) -> str:
    import docx
    doc = docx.Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    return "\n\n".join(parts)


def load_document(path: Path) -> List[Document]:
    """Загрузить один документ по пути (без импорта langchain_community → без PyTorch)."""
    suffix = path.suffix.lower()
    loaders = {
        ".txt": _load_txt,
        ".pdf": _load_pdf,
        ".docx": _load_docx,
    }
    if suffix not in loaders:
        raise ValueError(f"Неподдерживаемый формат: {suffix}. Доступны: .pdf, .txt, .docx")
    text = loaders[suffix](path)
    if not text.strip():
        return []
    return [Document(page_content=text.strip(), metadata={"source": str(path)})]


def load_documents_from_folder(
    folder: str | Path,
    glob: str = "**/*",
    suffixes: tuple = (".pdf", ".txt", ".docx"),
) -> List[Document]:
    """Загрузить все подходящие документы из папки."""
    folder = Path(folder)
    if not folder.is_dir():
        raise NotADirectoryError(str(folder))

    documents: List[Document] = []
    for path in folder.rglob("*"):
        if path.is_file() and path.suffix.lower() in suffixes:
            try:
                docs = load_document(path)
                documents.extend(docs)
            except Exception as e:
                print(f"Пропуск {path}: {e}")

    return documents


def _split_text_recursive(
    text: str,
    chunk_size: int,
    chunk_overlap: int,
    separators: List[str],
) -> List[str]:
    """Рекурсивное разбиение по разделителям без зависимости от langchain_text_splitters."""
    if not text or not text.strip():
        return []
    if len(text) <= chunk_size:
        return [text.strip()]

    sep = separators[0] if separators else ""
    if sep == "":
        # Разбиение по символам с перекрытием
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start = end - chunk_overlap
            if start >= len(text) - chunk_overlap:
                break
        if start < len(text):
            tail = text[start:].strip()
            if tail:
                chunks.append(tail)
        return chunks

    parts = text.split(sep)
    chunks: List[str] = []
    current = ""
    for i, part in enumerate(parts):
        piece = (sep + part) if i > 0 else part
        if len(current) + len(piece) <= chunk_size:
            current += piece
        else:
            if current.strip():
                chunks.append(current.strip())
            if len(piece) > chunk_size and len(separators) > 1:
                chunks.extend(
                    _split_text_recursive(
                        piece, chunk_size, chunk_overlap, separators[1:]
                    )
                )
                current = ""
            else:
                overlap_start = max(0, len(current) - chunk_overlap)
                current = current[overlap_start:] + piece
    if current.strip():
        chunks.append(current.strip())
    return chunks


def split_documents(
    documents: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    separators: List[str] | None = None,
) -> List[Document]:
    """Разбить документы на чанки (без зависимости от PyTorch/transformers)."""
    seps = separators or ["\n\n", "\n", ". ", " ", ""]
    result: List[Document] = []
    for doc in documents:
        texts = _split_text_recursive(
            doc.page_content, chunk_size, chunk_overlap, seps
        )
        for t in texts:
            if t.strip():
                result.append(Document(page_content=t.strip(), metadata=doc.metadata.copy()))
    return result
